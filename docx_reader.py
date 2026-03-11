# docx_reader.py — Extracts raw text from DOCX resumes

from docx import Document

def read_docx(file_path):
    
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
    
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""
    
    return text.strip()